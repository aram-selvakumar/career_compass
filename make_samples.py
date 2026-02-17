from docx import Document

def create_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience in Python and Web Development.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, Flask, JavaScript, HTML, CSS, SQL, Git, Docker')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Developed web applications using Flask and React.')
    doc.add_paragraph('Managed databases using PostgreSQL and MongoDB.')
    
    doc.save('test_resume.docx')
    print("Created test_resume.docx")

if __name__ == "__main__":
    create_resume()
